"""
DOCX Renderer for Report Export (v2.5.1 - FEAT-ARTIFACT-002)

Converts Markdown content to Microsoft Word (.docx) format using python-docx.

Features:
- Heading conversion (# -> Heading 1, ## -> Heading 2, etc.)
- Bold and italic text
- Bullet points and numbered lists
- Code blocks with monospace font
- Tables (when structured data detected)
- Configurable styling
- Extends FormatRenderer base class for unified interface

Usage (Legacy):
    renderer = DocxRenderer()
    docx_bytes = renderer.render(
        markdown_content="# Title\n\nContent...",
        title="Report Title",
        author="UBP System"
    )

Usage (v2.5.1 - Unified Interface):
    from renderers.base import RenderContext

    renderer = DocxRenderer()
    context = RenderContext(title="Report", content="# Title\n\nContent...")
    result = renderer.render(context)
"""

import io
import re
import logging
import time
from typing import Optional, List, Tuple, Dict, Any, Union
from dataclasses import dataclass
from enum import Enum

from .base import FormatRenderer, OutputFormat, RenderContext, RenderResult

logger = logging.getLogger(__name__)


class DocxStyle(str, Enum):
    """Available document styles."""
    PROFESSIONAL = "professional"
    TECHNICAL = "technical"
    MINIMAL = "minimal"


@dataclass
class StyleConfig:
    """Style configuration for document rendering."""
    title_font: str = "Calibri"
    body_font: str = "Calibri"
    code_font: str = "Consolas"
    title_size: int = 28  # pt
    heading1_size: int = 18
    heading2_size: int = 14
    heading3_size: int = 12
    body_size: int = 11
    code_size: int = 10
    line_spacing: float = 1.15


# Default styles
STYLE_CONFIGS = {
    DocxStyle.PROFESSIONAL: StyleConfig(
        title_font="Calibri Light",
        body_font="Calibri",
        code_font="Consolas",
    ),
    DocxStyle.TECHNICAL: StyleConfig(
        title_font="Arial",
        body_font="Arial",
        code_font="Courier New",
        line_spacing=1.0,
    ),
    DocxStyle.MINIMAL: StyleConfig(
        title_font="Helvetica",
        body_font="Helvetica",
        code_font="Monaco",
        title_size=24,
        heading1_size=16,
    ),
}


class DocxRenderer(FormatRenderer):
    """
    Renderer for converting Markdown to DOCX format.

    Extends FormatRenderer for unified interface while maintaining
    backward compatibility with legacy render() signature.

    Handles:
    - Headings (# ## ### ####)
    - Bold (**text**) and italic (*text*)
    - Bullet points (- or *)
    - Numbered lists (1. 2. 3.)
    - Code blocks (```)
    - Inline code (`code`)
    - Tables (| col1 | col2 |)
    """

    def __init__(
        self,
        style: DocxStyle = DocxStyle.PROFESSIONAL,
        settings: Optional[Dict[str, Any]] = None,
    ):
        """
        Initialize renderer with style.

        Args:
            style: Document style preset
            settings: Optional settings from ArtifactSettings
        """
        super().__init__(settings)
        self._style = style
        self._config = STYLE_CONFIGS.get(style, STYLE_CONFIGS[DocxStyle.PROFESSIONAL])
        self._logger.info(f"[DOCX] Renderer initialized with style: {style}")

    @property
    def output_format(self) -> OutputFormat:
        """Return DOCX output format."""
        return OutputFormat.DOCX

    def render(
        self,
        content_or_context: Union[str, RenderContext],
        title: str = "Report",
        author: str = "UBP Enterprise",
        subject: str = "",
        metadata: Optional[Dict[str, Any]] = None,
    ) -> Union[bytes, RenderResult]:
        """
        Render Markdown content to DOCX.

        Supports two interfaces:
        1. Legacy: render(markdown_content, title, author, ...) -> bytes
        2. Unified: render(RenderContext) -> RenderResult

        Args:
            content_or_context: Either markdown string (legacy) or RenderContext (unified)
            title: Document title (legacy interface)
            author: Document author (legacy interface)
            subject: Document subject (legacy interface)
            metadata: Additional metadata (legacy interface)

        Returns:
            bytes (legacy) or RenderResult (unified)
        """
        # Detect interface type
        if isinstance(content_or_context, RenderContext):
            return self._render_unified(content_or_context)
        else:
            # Legacy interface - return bytes directly
            return self._render_legacy(
                content_or_context, title, author, subject, metadata
            )

    def _render_unified(self, context: RenderContext) -> RenderResult:
        """Render using unified RenderContext interface."""
        start_time = time.time()
        self._log_render_start(context)

        try:
            docx_bytes = self._render_legacy(
                markdown_content=context.content,
                title=context.title,
                author=context.author,
                subject=context.subject,
                metadata=context.metadata,
            )

            duration_ms = int((time.time() - start_time) * 1000)
            filename = self.generate_filename(context.title, context.blueprint_id)

            result = RenderResult(
                content=docx_bytes,
                format=self.output_format,
                filename=filename,
                render_time_ms=duration_ms,
                metadata={
                    "title": context.title,
                    "author": context.author,
                    "content_length": len(context.content),
                },
            )

            self._log_render_complete(result, duration_ms)
            return result

        except Exception as e:
            self._logger.error(f"[DOCX] Render error: {e}", exc_info=True)
            return RenderResult.failure(
                format=self.output_format,
                error=str(e),
                filename=self.generate_filename(context.title, context.blueprint_id),
            )

    def _render_legacy(
        self,
        markdown_content: str,
        title: str = "Report",
        author: str = "UBP Enterprise",
        subject: str = "",
        metadata: Optional[Dict[str, Any]] = None,
    ) -> bytes:
        """
        Legacy render method - returns bytes directly.

        Args:
            markdown_content: Markdown-formatted text
            title: Document title
            author: Document author
            subject: Document subject
            metadata: Additional metadata

        Returns:
            DOCX file as bytes
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
        except ImportError:
            logger.error("[DOCX] python-docx not installed")
            raise ImportError("python-docx is required for DOCX rendering")

        # Create document
        doc = Document()

        # Set document properties
        core_props = doc.core_properties
        core_props.title = title
        core_props.author = author
        core_props.subject = subject or f"Generated Report: {title}"

        # Add title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata if provided
        if metadata:
            meta_para = doc.add_paragraph()
            meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if "generated_at" in metadata:
                meta_run = meta_para.add_run(f"Generated: {metadata['generated_at']}")
                meta_run.font.size = Pt(9)
                meta_run.font.italic = True

        # Add a line break after title
        doc.add_paragraph()

        # Parse and render markdown
        self._render_markdown(doc, markdown_content)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        docx_bytes = buffer.getvalue()
        logger.info(f"[DOCX] Rendered document: {len(docx_bytes)} bytes")

        return docx_bytes

    def _render_markdown(self, doc, content: str) -> None:
        """Parse and render markdown content to document."""
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        lines = content.split('\n')
        i = 0
        in_code_block = False
        code_block_lines = []
        in_table = False
        table_rows = []

        while i < len(lines):
            line = lines[i]

            # Handle code blocks
            if line.strip().startswith('```'):
                if in_code_block:
                    # End code block
                    self._add_code_block(doc, code_block_lines)
                    code_block_lines = []
                    in_code_block = False
                else:
                    # Start code block
                    in_code_block = True
                i += 1
                continue

            if in_code_block:
                code_block_lines.append(line)
                i += 1
                continue

            # Handle tables
            if '|' in line and line.strip().startswith('|'):
                if not in_table:
                    in_table = True
                    table_rows = []

                # Skip separator rows (|---|---|)
                if not re.match(r'^\s*\|[\s\-:|]+\|\s*$', line):
                    table_rows.append(line)
                i += 1
                continue
            elif in_table:
                # End of table
                self._add_table(doc, table_rows)
                table_rows = []
                in_table = False
                # Don't increment i, process current line

            # Handle headings
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2).strip()
                doc.add_heading(text, level=min(level, 4))
                i += 1
                continue

            # Handle bullet points
            bullet_match = re.match(r'^[\s]*[-*]\s+(.+)$', line)
            if bullet_match:
                text = bullet_match.group(1)
                para = doc.add_paragraph(style='List Bullet')
                self._add_formatted_text(para, text)
                i += 1
                continue

            # Handle numbered lists
            number_match = re.match(r'^[\s]*(\d+)\.\s+(.+)$', line)
            if number_match:
                text = number_match.group(2)
                para = doc.add_paragraph(style='List Number')
                self._add_formatted_text(para, text)
                i += 1
                continue

            # Handle regular paragraphs
            if line.strip():
                para = doc.add_paragraph()
                self._add_formatted_text(para, line)

            i += 1

        # Handle any remaining table
        if in_table and table_rows:
            self._add_table(doc, table_rows)

    def _add_formatted_text(self, paragraph, text: str) -> None:
        """Add text with inline formatting (bold, italic, code)."""
        from docx.shared import Pt, RGBColor

        # Pattern for bold, italic, and code
        # Process in order: code first (to avoid conflicts), then bold, then italic
        pattern = r'(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)'

        parts = re.split(pattern, text)

        for part in parts:
            if not part:
                continue

            if part.startswith('`') and part.endswith('`'):
                # Inline code
                code_text = part[1:-1]
                run = paragraph.add_run(code_text)
                run.font.name = self._config.code_font
                run.font.size = Pt(self._config.code_size)
                # Light gray background effect (approximate)
                run.font.color.rgb = RGBColor(80, 80, 80)
            elif part.startswith('**') and part.endswith('**'):
                # Bold
                bold_text = part[2:-2]
                run = paragraph.add_run(bold_text)
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                # Italic
                italic_text = part[1:-1]
                run = paragraph.add_run(italic_text)
                run.italic = True
            else:
                # Regular text
                paragraph.add_run(part)

    def _add_code_block(self, doc, lines: List[str]) -> None:
        """Add a code block with monospace font and background."""
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        if not lines:
            return

        # Add code as a single paragraph with special formatting
        para = doc.add_paragraph()

        # Set paragraph formatting for code block
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.right_indent = Inches(0.5)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)

        code_text = '\n'.join(lines)
        run = para.add_run(code_text)
        run.font.name = self._config.code_font
        run.font.size = Pt(self._config.code_size)
        run.font.color.rgb = RGBColor(40, 40, 40)

    def _add_table(self, doc, rows: List[str]) -> None:
        """Add a table from markdown table rows."""
        from docx.shared import Pt, Inches

        if not rows:
            return

        # Parse table structure
        parsed_rows = []
        for row in rows:
            # Split by | and clean up
            cells = [cell.strip() for cell in row.split('|')]
            # Remove empty first/last cells from | borders
            cells = [c for c in cells if c]
            if cells:
                parsed_rows.append(cells)

        if not parsed_rows:
            return

        # Determine number of columns
        num_cols = max(len(row) for row in parsed_rows)

        # Create table
        table = doc.add_table(rows=len(parsed_rows), cols=num_cols)
        table.style = 'Table Grid'

        # Populate cells
        for i, row_data in enumerate(parsed_rows):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                if j < num_cols:
                    cell = row.cells[j]
                    # First row as header (bold)
                    if i == 0:
                        run = cell.paragraphs[0].add_run(cell_text)
                        run.bold = True
                    else:
                        cell.text = cell_text

        # Add space after table
        doc.add_paragraph()

        logger.debug(f"[DOCX] Added table: {len(parsed_rows)} rows x {num_cols} cols")


class MarkdownRenderer(FormatRenderer):
    """
    Simple renderer that outputs clean Markdown.

    Useful for exporting reports as .md files without conversion.
    Extends FormatRenderer for unified interface.
    """

    def __init__(self, settings: Optional[Dict[str, Any]] = None):
        """Initialize Markdown renderer."""
        super().__init__(settings)

    @property
    def output_format(self) -> OutputFormat:
        """Return Markdown output format."""
        return OutputFormat.MD

    def render(
        self,
        content_or_context: Union[str, RenderContext],
        title: str = "Report",
        metadata: Optional[Dict[str, Any]] = None,
    ) -> Union[bytes, RenderResult]:
        """
        Render content as Markdown.

        Supports both legacy and unified interfaces.

        Args:
            content_or_context: Either markdown string (legacy) or RenderContext (unified)
            title: Document title (legacy interface)
            metadata: Optional metadata (legacy interface)

        Returns:
            bytes (legacy) or RenderResult (unified)
        """
        if isinstance(content_or_context, RenderContext):
            return self._render_unified(content_or_context)
        else:
            return self._render_legacy(content_or_context, title, metadata)

    def _render_unified(self, context: RenderContext) -> RenderResult:
        """Render using unified RenderContext interface."""
        start_time = time.time()
        self._log_render_start(context)

        try:
            md_bytes = self._render_legacy(
                content=context.content,
                title=context.title,
                metadata=context.metadata,
            )

            duration_ms = int((time.time() - start_time) * 1000)
            filename = self.generate_filename(context.title, context.blueprint_id)

            result = RenderResult(
                content=md_bytes,
                format=self.output_format,
                filename=filename,
                render_time_ms=duration_ms,
                metadata={
                    "title": context.title,
                    "content_length": len(context.content),
                },
            )

            self._log_render_complete(result, duration_ms)
            return result

        except Exception as e:
            self._logger.error(f"[MD] Render error: {e}", exc_info=True)
            return RenderResult.failure(
                format=self.output_format,
                error=str(e),
                filename=self.generate_filename(context.title, context.blueprint_id),
            )

    def _render_legacy(
        self,
        content: str,
        title: str = "Report",
        metadata: Optional[Dict[str, Any]] = None,
    ) -> bytes:
        """
        Legacy render method - returns bytes directly.

        Args:
            content: Markdown content
            title: Document title (prepended as H1)
            metadata: Optional metadata to include

        Returns:
            UTF-8 encoded Markdown bytes
        """
        output_lines = []

        # Add title
        output_lines.append(f"# {title}")
        output_lines.append("")

        # Add metadata as YAML frontmatter
        if metadata:
            output_lines.append("---")
            for key, value in metadata.items():
                output_lines.append(f"{key}: {value}")
            output_lines.append("---")
            output_lines.append("")

        # Add content
        output_lines.append(content)

        output = '\n'.join(output_lines)
        return output.encode('utf-8')
